"""Text extraction from PDF, DOCX and TXT."""

import io

import pytest
from docx import Document as DocxDocument

from app.core.exceptions import EmptyDocumentError, TextExtractionError, UnsupportedFileTypeError
from app.services.text_extraction import extract_text, normalise, truncate
from tests.fixtures.pdf_builder import build_pdf, build_pdf_without_text_layer


def build_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    document = DocxDocument()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        added = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for cell_index, value in enumerate(row):
                added.cell(row_index, cell_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestTxt:
    def test_utf8(self) -> None:
        assert extract_text(b"Invoice total: 42.00\n", ".txt") == "Invoice total: 42.00"

    def test_bom_stripped(self) -> None:
        assert extract_text("Hello".encode("utf-8-sig"), ".txt") == "Hello"

    def test_cp1252_fallback(self) -> None:
        """Text exported from older Windows tooling still decodes."""
        assert "Ré" in extract_text("Réçu".encode("cp1252"), ".txt")

    def test_undecodable_is_permanent(self) -> None:
        with pytest.raises(TextExtractionError) as error:
            extract_text(b"\xff\xfe\x00\x80\x81\x8d", ".txt")
        assert error.value.retryable is False

    def test_whitespace_only_is_empty(self) -> None:
        with pytest.raises(EmptyDocumentError):
            extract_text(b"   \n\n\t  \n", ".txt")


class TestDocx:
    def test_paragraphs(self) -> None:
        text = extract_text(build_docx(["First line", "Second line"]), ".docx")
        assert "First line" in text
        assert "Second line" in text

    def test_table_cells_included(self) -> None:
        """Invoices put their line items in tables; missing them loses the data."""
        text = extract_text(
            build_docx(["Invoice"], table=[["Item", "Amount"], ["Widget", "42.00"]]), ".docx"
        )
        assert "Widget" in text
        assert "42.00" in text

    def test_empty_document_is_permanent(self) -> None:
        with pytest.raises(EmptyDocumentError):
            extract_text(build_docx([]), ".docx")

    def test_corrupt_file_is_permanent(self) -> None:
        with pytest.raises((TextExtractionError, EmptyDocumentError)) as error:
            extract_text(b"PK\x03\x04" + b"\x00" * 100, ".docx")
        assert error.value.retryable is False


class TestPdf:
    def test_text_layer_extracted(self, pdf_bytes: bytes) -> None:
        assert "Sample invoice" in extract_text(pdf_bytes, ".pdf")

    def test_multiline_text_preserved(self) -> None:
        text = extract_text(build_pdf(["First line", "Second line"]), ".pdf")
        assert "First line" in text
        assert "Second line" in text

    def test_pdf_without_a_text_layer_fails_permanently(self) -> None:
        """A scanned document needs OCR, which this pipeline does not do; the
        job must fail immediately rather than burn its retry budget."""
        with pytest.raises(EmptyDocumentError) as error:
            extract_text(build_pdf_without_text_layer(), ".pdf")
        assert error.value.retryable is False

    def test_corrupt_pdf_is_permanent(self) -> None:
        with pytest.raises((TextExtractionError, EmptyDocumentError)) as error:
            extract_text(b"%PDF-1.4\nthis is not really a pdf", ".pdf")
        assert error.value.retryable is False


def test_unknown_extension_rejected() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        extract_text(b"anything", ".rtf")


class TestNormalise:
    def test_blank_line_runs_collapse(self) -> None:
        assert normalise("a\n\n\n\n\nb") == "a\n\nb"

    def test_trailing_whitespace_removed(self) -> None:
        assert normalise("a   \nb\t\n") == "a\nb"

    def test_line_endings_unified(self) -> None:
        assert normalise("a\r\nb\rc") == "a\nb\nc"

    def test_leading_and_trailing_blank_lines_removed(self) -> None:
        assert normalise("\n\n  content  \n\n") == "content"


class TestTruncate:
    def test_short_text_untouched(self) -> None:
        assert truncate("hello", 100) == ("hello", False)

    def test_long_text_clipped_and_flagged(self) -> None:
        text, was_truncated = truncate("x" * 500, 100)
        assert len(text) == 100
        assert was_truncated is True

    def test_exact_length_not_truncated(self) -> None:
        assert truncate("x" * 100, 100) == ("x" * 100, False)
