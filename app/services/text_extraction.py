"""Turning uploaded bytes into plain text.

Every failure in here is permanent. A PDF that will not parse now will not
parse in thirty seconds either, so these raise non-retryable errors and the job
goes straight to ``failed`` with a message explaining why, rather than burning
three attempts first.

``pypdf`` is used rather than PyMuPDF: PyMuPDF is AGPL-licensed, which is a
real consideration for a public repository, and its speed advantage is
irrelevant at this scale.
"""

import io
import zipfile
from collections.abc import Callable

import pypdf
from docx import Document as DocxDocument
from docx.opc.exceptions import OpcError
from pypdf.errors import PyPdfError

from app.core.exceptions import EmptyDocumentError, TextExtractionError, UnsupportedFileTypeError
from app.core.logging import get_logger

logger = get_logger(__name__)

#: Encodings tried in order for .txt uploads. ``utf-8-sig`` comes first because
#: it decodes plain UTF-8 identically *and* strips a byte-order mark; plain
#: ``utf-8`` would succeed on a BOM'd file and leave U+FEFF in the text.
#: cp1252 catches text exported from older Windows tooling.
TEXT_ENCODINGS = ("utf-8-sig", "cp1252")


def extract_pdf(data: bytes) -> str:
    """Extract the text layer of a PDF.

    Scanned documents have no text layer and produce an empty string, which the
    caller turns into a permanent failure. OCR is out of scope.
    """
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except (PyPdfError, ValueError, OSError) as error:
        raise TextExtractionError(
            "The PDF could not be parsed.", details={"format": "pdf"}
        ) from error
    return "\n\n".join(pages)


def extract_docx(data: bytes) -> str:
    """Extract paragraph and table text from a DOCX file."""
    try:
        document = DocxDocument(io.BytesIO(data))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        blocks.extend(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            # Word repeats a merged cell's text once per spanned column.
            if cell.text
        )
    except (OpcError, zipfile.BadZipFile, ValueError, KeyError, OSError) as error:
        # A .docx is a zip; a truncated or mislabelled one surfaces as a zip
        # error or an OPC package error rather than anything docx-specific.
        raise TextExtractionError(
            "The DOCX file could not be parsed.", details={"format": "docx"}
        ) from error
    return "\n".join(blocks)


def extract_txt(data: bytes) -> str:
    """Decode a plain-text upload, tolerating a couple of common encodings."""
    for encoding in TEXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise TextExtractionError(
        "The text file is not in a supported encoding.",
        details={"tried": list(TEXT_ENCODINGS)},
    )


EXTRACTORS: dict[str, Callable[[bytes], str]] = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_txt,
}


def normalise(text: str) -> str:
    """Collapse the whitespace noise that PDF extraction tends to produce.

    Blank-line runs become a single separator and trailing spaces are dropped,
    which measurably reduces the token count sent to the model without losing
    structure the model relies on.
    """
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]

    collapsed: list[str] = []
    blank_run = 0
    for line in lines:
        if line:
            blank_run = 0
            collapsed.append(line)
        else:
            blank_run += 1
            if blank_run == 1:
                collapsed.append("")

    return "\n".join(collapsed).strip()


def extract_text(data: bytes, extension: str) -> str:
    """Extract normalised text from an uploaded file.

    Raises:
        UnsupportedFileTypeError: no extractor exists for this extension.
        TextExtractionError: the file could not be parsed.
        EmptyDocumentError: parsing succeeded but yielded no usable text.
    """
    extractor = EXTRACTORS.get(extension.lower())
    if extractor is None:
        raise UnsupportedFileTypeError(
            f"No text extractor is registered for {extension}.",
            details={"extension": extension},
        )

    text = normalise(extractor(data))
    if not text:
        raise EmptyDocumentError(
            "The document contains no extractable text. Scanned or image-only "
            "documents require OCR, which this pipeline does not perform.",
            details={"extension": extension},
        )

    logger.info("text.extracted", extension=extension, characters=len(text))
    return text


def truncate(text: str, max_chars: int) -> tuple[str, bool]:
    """Clip text to a character budget, reporting whether anything was cut."""
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True
